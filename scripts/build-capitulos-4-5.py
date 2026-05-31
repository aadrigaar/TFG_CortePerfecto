from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "entregas"
DIAGRAM_DIR = ROOT / "diagramas" / "capitulo4" / "imagenes"
SCREENSHOT_DIR = ROOT / "diagramas" / "capitulo4" / "capturas"
DOCX_PATH = OUT_DIR / "Capitulos4y5.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "C9A13A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TEXT = "1F2937"
MUTED = "56616F"
BORDER = "D9E1EA"


def font_path() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/segoeui.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if bold:
        candidates = [
            Path("C:/Windows/Fonts/calibrib.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            Path("C:/Windows/Fonts/segoeuib.ttf"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return ImageFont.truetype(str(candidate), size)

    path = font_path()
    return ImageFont.truetype(path, size) if path else ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str, text: str, subtitle: str = "") -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = box
    title_font = get_font(26, bold=True)
    body_font = get_font(18)
    lines = wrap_text(draw, text, title_font, x2 - x1 - 36)
    y = y1 + 22
    for line in lines:
        draw.text((x1 + 18, y), line, font=title_font, fill="#101828")
        y += 30
    if subtitle:
        for line in wrap_text(draw, subtitle, body_font, x2 - x1 - 36):
            draw.text((x1 + 18, y + 8), line, font=body_font, fill="#475467")
            y += 24


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#667085") -> None:
    draw.line((start, end), fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if abs(ey - sy) > abs(ex - sx):
        if ey >= sy:
            points = [(ex, ey), (ex - 8, ey - 14), (ex + 8, ey - 14)]
        else:
            points = [(ex, ey), (ex - 8, ey + 14), (ex + 8, ey + 14)]
    elif ex >= sx:
        points = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
    else:
        points = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    draw.polygon(points, fill=color)


def create_navigation_map() -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / "figura_4_1_mapa_navegacion.png"
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(34, bold=True)
    subtitle_font = get_font(20)

    draw.text((70, 45), "Mapa navegable de la solución implementada", font=title_font, fill="#0B2545")
    draw.text(
        (70, 90),
        "El cliente accede por la web pública y el peluquero por el panel privado protegido con JWT.",
        font=subtitle_font,
        fill="#475467",
    )

    public = (70, 170, 830, 890)
    admin = (970, 170, 1730, 890)
    draw.rounded_rectangle(public, radius=24, fill="#F7FAFC", outline="#BFD7F2", width=3)
    draw.rounded_rectangle(admin, radius=24, fill="#FFFCF2", outline="#E8C75F", width=3)
    draw.text((105, 205), "Cliente", font=get_font(28, bold=True), fill="#1F4D78")
    draw.text((1005, 205), "Administrador / peluquero", font=get_font(28, bold=True), fill="#7A5A00")

    public_boxes = [
        ((130, 285, 380, 385), "Inicio", "Propuesta de valor y CTA"),
        ((500, 285, 760, 385), "Servicios", "Catálogo y precios"),
        ((130, 470, 380, 570), "Chatbot IA", "Reserva guiada"),
        ((500, 470, 760, 570), "Confirmación", "Tarjeta de cita"),
        ((315, 650, 585, 760), "Contacto", "Ubicación y horario"),
    ]
    for box, text, subtitle in public_boxes:
        rounded_box(draw, box, "#FFFFFF", "#BFD7F2", text, subtitle)
    arrow(draw, (380, 335), (500, 335))
    arrow(draw, (255, 385), (255, 470))
    arrow(draw, (380, 520), (500, 520))
    arrow(draw, (630, 570), (510, 650))

    admin_boxes = [
        ((1030, 285, 1285, 385), "Login", "Credenciales admin"),
        ((1410, 285, 1670, 385), "Dashboard", "KPIs y próximas citas"),
        ((1030, 470, 1285, 570), "Gestión de citas", "Filtrar, editar, completar"),
        ((1410, 470, 1670, 570), "Crear cita", "Alta manual"),
        ((1220, 650, 1485, 760), "MongoDB", "Persistencia"),
    ]
    for box, text, subtitle in admin_boxes:
        rounded_box(draw, box, "#FFFFFF", "#E8C75F", text, subtitle)
    arrow(draw, (1285, 335), (1410, 335), "#8A6A00")
    arrow(draw, (1540, 385), (1180, 470), "#8A6A00")
    arrow(draw, (1285, 520), (1410, 520), "#8A6A00")
    arrow(draw, (1175, 570), (1320, 650), "#8A6A00")
    arrow(draw, (1540, 570), (1395, 650), "#8A6A00")

    draw.line((900, 220, 900, 845), fill="#D0D5DD", width=3)
    draw.text((790, 920), "API REST Node.js/Express + servicios de negocio + Mongoose", font=get_font(24, bold=True), fill="#0B2545")

    img.save(path)
    return path


def create_context_navigation_diagram() -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / "figura_4_1_diagrama_contexto_navegacion.png"
    img = Image.new("RGB", (1900, 1200), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "Diagrama de contexto y navegación", font=get_font(34, bold=True), fill="#0B2545")
    draw.text(
        (70, 90),
        "El orden lógico viene de los actores y la navegación se estructura por casos de uso.",
        font=get_font(20),
        fill="#475467",
    )

    # Actors
    rounded_box(draw, (70, 250, 310, 360), "#F7FAFC", "#98C1E6", "Cliente", "Usuario público")
    rounded_box(draw, (70, 760, 310, 870), "#FFFCF2", "#C9A13A", "Peluquero", "Administrador")

    # Public navigation
    rounded_box(draw, (390, 185, 700, 315), "#FFFFFF", "#98C1E6", "Web pública", "Inicio, servicios, combos, nosotros, opiniones y contacto")
    rounded_box(draw, (790, 185, 1100, 315), "#FFFFFF", "#98C1E6", "Chatbot IA", "Conversación de reserva y consultas")
    rounded_box(draw, (1190, 185, 1500, 315), "#FFFFFF", "#98C1E6", "Cita confirmada", "Tarjeta visible al cliente")

    # Admin navigation
    rounded_box(draw, (390, 700, 700, 830), "#FFFFFF", "#C9A13A", "Login admin", "Validación de credenciales")
    rounded_box(draw, (790, 675, 1100, 825), "#FFFFFF", "#C9A13A", "Panel admin", "Dashboard, agenda y creación manual")
    rounded_box(draw, (1190, 675, 1500, 825), "#FFFFFF", "#C9A13A", "Gestión de citas", "Filtrar, ordenar, editar, completar y eliminar")

    # Common technical context
    rounded_box(draw, (1560, 385, 1850, 525), "#F7FAFC", "#1F4D78", "API Node.js/Express", "Controladores, servicios y validaciones")
    rounded_box(draw, (1560, 115, 1850, 235), "#FFFCF2", "#C9A13A", "LM Studio local", "Meta Llama 3.1 8B")
    rounded_box(draw, (1560, 635, 1850, 755), "#F4F6F9", "#667085", "MongoDB", "appointments, admins y servicios")

    # Arrows: public flow
    arrow(draw, (310, 305), (390, 250), "#475467")
    arrow(draw, (700, 250), (790, 250), "#475467")
    arrow(draw, (1100, 250), (1190, 250), "#475467")
    arrow(draw, (1100, 285), (1560, 435), "#475467")
    arrow(draw, (1560, 425), (1500, 285), "#475467")

    # Arrows: admin flow
    arrow(draw, (310, 815), (390, 765), "#8A6A00")
    arrow(draw, (700, 765), (790, 750), "#8A6A00")
    arrow(draw, (1100, 750), (1190, 750), "#8A6A00")
    arrow(draw, (1500, 750), (1560, 475), "#8A6A00")

    # Backend integrations
    arrow(draw, (1705, 385), (1705, 235), "#8A6A00")
    arrow(draw, (1785, 235), (1785, 385), "#8A6A00")
    arrow(draw, (1705, 525), (1705, 635), "#475467")

    # Internal labels
    label_font = get_font(18)
    labels = [
        (440, 340, "UC-01/02/03"),
        (835, 340, "UC-04/05/06/07/08/09"),
        (440, 855, "UC-10"),
        (835, 855, "UC-11/12/13"),
        (1230, 855, "UC-14/15/16/17"),
    ]
    for x, y, label in labels:
        draw.rounded_rectangle((x, y, x + 210, y + 42), radius=14, fill="#EEF4FF", outline="#BFD7F2", width=1)
        draw.text((x + 18, y + 10), label, font=label_font, fill="#1F4D78")

    # Rule box
    draw.rounded_rectangle((430, 930, 850, 1135), radius=18, fill="#F7FAFC", outline="#D0D5DD", width=2)
    draw.text((455, 955), "Reglas críticas", font=get_font(24, bold=True), fill="#1F4D78")
    for idx, text in enumerate(["Servicio 1..7", "Lunes a viernes", "10:00-20:00", "Sin solapes", "Hora futura"]):
        draw.text((455, 995 + idx * 25), f"- {text}", font=get_font(18), fill="#344054")

    draw.rounded_rectangle((1130, 955, 1480, 1115), radius=18, fill="#FFFCF2", outline="#E8C75F", width=2)
    draw.text((1155, 980), "Cierre de ciclo", font=get_font(24, bold=True), fill="#7A5A00")
    for idx, text in enumerate(["Confirmada", "Completada", "Cancelada", "Eliminada"]):
        draw.text((1155, 1018 + idx * 25), f"- {text}", font=get_font(18), fill="#344054")

    img.save(path)
    return path


def create_booking_flow() -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / "figura_4_9_flujo_reserva.png"
    img = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "Flujo implementado de reserva por chatbot", font=get_font(34, bold=True), fill="#0B2545")
    draw.text((70, 90), "La IA conversa, pero las reglas críticas se validan siempre en backend.", font=get_font(20), fill="#475467")

    boxes = [
        ((70, 220, 310, 340), "Cliente", "Mensaje natural"),
        ((380, 220, 620, 340), "ChatWidget", "Historial y tarjeta"),
        ((690, 220, 930, 340), "chatController", "Orquesta el flujo"),
        ((1030, 110, 1300, 230), "bookingFlowService", "Reserva determinista"),
        ((1030, 270, 1300, 390), "chatRuleService", "Reglas previas"),
        ((1390, 110, 1660, 230), "prompt + LM Studio", "Conversación natural"),
        ((1390, 270, 1660, 390), "responseParserService", "Respuesta y candidato"),
        ((1030, 455, 1300, 575), "appointmentService", "Validación final"),
        ((1390, 455, 1660, 575), "MongoDB", "Cita persistida"),
    ]
    for box, text, subtitle in boxes:
        fill = "#FFFCF2" if "LM Studio" in text else "#F7FAFC"
        outline = "#E8C75F" if "LM Studio" in text else "#BFD7F2"
        rounded_box(draw, box, fill, outline, text, subtitle)
    arrow(draw, (310, 280), (380, 280))
    arrow(draw, (620, 280), (690, 280))
    arrow(draw, (930, 245), (1030, 170))
    arrow(draw, (930, 305), (1030, 330))
    arrow(draw, (1300, 170), (1390, 170))
    arrow(draw, (1525, 230), (1525, 270))
    arrow(draw, (1165, 230), (1165, 455))
    arrow(draw, (1390, 330), (1300, 515))
    arrow(draw, (1300, 515), (1390, 515))

    checks = [
        "Fin de semana bloqueado antes de pedir hora",
        "Hora pasada o fuera de 10:00-20:00 rechazada",
        "Nombre inválido no se persiste",
        "Solapes detectados contra citas activas",
        "Catálogo 1..7 como fuente única de verdad",
    ]
    x, y = 220, 650
    draw.rounded_rectangle((180, 610, 1620, 840), radius=22, fill="#F4F6F9", outline="#D0D5DD", width=2)
    draw.text((220, 640), "Validaciones deterministas", font=get_font(26, bold=True), fill="#1F4D78")
    for i, item in enumerate(checks):
        draw.text((x + (i % 2) * 680, y + 42 + (i // 2) * 42), f"- {item}", font=get_font(21), fill="#344054")

    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx] / 567)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "TFG Corte Perfecto · Capítulos 4 y 5"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Adrián García Arranz · Ingeniería Informática"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("CAPÍTULOS 4 Y 5")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Implementación, validación, conclusiones y líneas futuras")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Corte Perfecto · Peluquería en Santander con reservas mediante IA local")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Trabajo de Fin de Grado · Ingeniería Informática")
    run.font.size = Pt(12)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autor: Adrián García Arranz")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Curso 2025/2026")
    run.font.size = Pt(12)

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], LIGHT_BLUE)
        set_cell_margins(header_cells[i])
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in header_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string("0B2545")

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value
            set_cell_margins(row_cells[i])
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in row_cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor.from_string(TEXT)

    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    set_table_width(table, [9000])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def build_document() -> None:
    context_diagram = create_context_navigation_diagram()
    nav_map = create_navigation_map()
    booking_flow = create_booking_flow()

    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_page(doc)

    doc.add_heading("Capítulo 4. Implementación y mapa de la solución", level=1)
    doc.add_paragraph(
        "Este capítulo presenta la solución desarrollada desde el punto de vista de su materialización final. "
        "Después de haber definido requisitos, análisis y diseño en los capítulos anteriores, se muestra cómo esos "
        "artefactos se convierten en pantallas, flujos navegables, rutas de API, servicios de negocio y colecciones "
        "persistidas en MongoDB. El objetivo no es repetir el diseño técnico, sino demostrar que existe una aplicación "
        "funcional y coherente con el proceso metodológico seguido."
    )

    add_callout(
        doc,
        "Criterio de implementación",
        "La solución se considera cerrada cuando cada caso de uso principal puede recorrerse desde la interfaz, pasar por la API Express, aplicar reglas de negocio en servicios independientes y dejar evidencia persistente o verificable.",
    )

    doc.add_heading("4.1 Estado técnico de la solución implementada", level=2)
    doc.add_paragraph(
        "Antes de documentar la interfaz se revisó el estado del código para comprobar que la implementación respeta "
        "la arquitectura comprometida: frontend React/Vite, backend Node.js/Express organizado siguiendo MVC, "
        "persistencia MongoDB mediante Mongoose y conexión local con LM Studio. La comprobación ejecutada fue "
        "`npm run verify`, que encadena validación sintáctica del backend, pruebas automáticas y construcción de "
        "producción del frontend."
    )

    add_table(
        doc,
        ["Área revisada", "Resultado", "Evidencia"],
        [
            ["Estructura MVC backend", "Rutas, controladores, servicios y modelos separados.", "backend/src/routes, controllers, services, models"],
            ["Persistencia MongoDB", "Colecciones `appointments`, `admins` y `servicios` modeladas o sincronizadas.", "Appointment.js, Admin.js, Service.js"],
            ["Chatbot con IA local", "LM Studio actúa como motor conversacional; las reglas críticas permanecen en backend.", "chatController, bookingFlowService, lmStudioService"],
            ["Panel administrador", "Login JWT, dashboard, listado, filtrado, creación, edición, completado y borrado.", "AdminLogin, AdminDashboard, AdminAppointments, AppointmentForm"],
            ["Verificación automática", "10 pruebas superadas, backend válido y build de frontend correcto.", "backend/tests y `npm run verify`"],
        ],
        [2100, 4100, 3160],
    )

    doc.add_heading("4.2 Organización final del código", level=2)
    doc.add_paragraph(
        "La organización del proyecto se mantiene deliberadamente simple. No se introducen microservicios ni capas "
        "accidentales porque el alcance del TFG se resuelve mejor con una API modular de alta cohesión. La separación "
        "entre frontend y backend es física y lógica: el frontend no accede a MongoDB ni a LM Studio; todo pasa por la "
        "API. A su vez, los controladores no concentran reglas de negocio, sino que delegan en servicios."
    )

    add_table(
        doc,
        ["Capa", "Carpetas / módulos", "Responsabilidad"],
        [
            ["Presentación pública", "frontend/src/pages, components, styles", "Mostrar la web comercial, abrir el chatbot y renderizar la confirmación de cita."],
            ["Presentación administrativa", "frontend/src/pages/admin, components/admin, context", "Login, dashboard, filtros de agenda, formularios y acciones privadas."],
            ["Entrada backend", "backend/src/routes, controllers, middleware", "Exponer API REST, proteger rutas privadas, controlar errores y coordinar cada petición."],
            ["Dominio / negocio", "backend/src/services, config/serviceCatalog.js", "Gestionar citas, autenticación, calendario, solapes, servicios, prompt, reglas de chat y respuesta de LM Studio."],
            ["Persistencia", "backend/src/models, MongoDB", "Persistir citas, cuenta administrativa y catálogo de servicios sincronizado."],
            ["Evidencia metodológica", "RUP/99-seguimiento, backend/tests", "Mantener trazabilidad caso de uso-código-prueba siguiendo el estilo de pySigHor."],
        ],
        [1800, 3000, 4560],
    )

    doc.add_paragraph(
        "Durante la revisión final se eliminaron responsabilidades mezcladas: la sincronización del catálogo de servicios "
        "queda centralizada en `serviceCatalogService` y la preparación/autenticación de la cuenta del administrador se "
        "ubica en `adminService`. Por tanto, `database.js` conserva una responsabilidad única: abrir la conexión con "
        "MongoDB. Esta decisión mejora la cohesión y evita que el módulo de conexión conozca reglas funcionales."
    )

    doc.add_heading("4.3 Mapa navegable de la solución", level=2)
    doc.add_paragraph(
        "El mapa de navegación se deriva de los actores y casos de uso definidos en el Capítulo 2. El cliente entra en "
        "la web pública y puede consultar información o reservar mediante el chatbot. El administrador accede por una "
        "ruta privada, se autentica y gestiona la agenda. Ambas experiencias convergen en la misma API y en la misma "
        "base de datos, lo que evita agendas paralelas."
    )
    add_picture(doc, context_diagram, "Figura 4.1. Diagrama de contexto y navegación de Corte Perfecto.")
    doc.add_paragraph(
        "Este diagrama cumple la indicación de partir del diagrama de contexto: no representa solo pantallas, sino "
        "transiciones relevantes entre actores, sistema público, chatbot, panel privado y persistencia. La navegación "
        "queda ligada a los casos de uso, por lo que el Capítulo 4 no se limita a enseñar interfaces aisladas."
    )
    add_picture(doc, nav_map, "Figura 4.2. Mapa navegable de la solución implementada.")

    add_table(
        doc,
        ["Zona", "Pantallas", "Casos de uso cubiertos"],
        [
            ["Web pública", "Inicio, Servicios, Combos, Nosotros, Opiniones, Contacto", "UC-01, UC-02, UC-03"],
            ["Chatbot", "Widget flotante, historial, acciones rápidas, tarjeta de cita", "UC-04, UC-05, UC-06, UC-07, UC-08, UC-09"],
            ["Acceso admin", "Login protegido", "UC-10"],
            ["Panel admin", "Dashboard, Gestión de Citas, Crear Cita, Modal de edición", "UC-11, UC-12, UC-13, UC-14, UC-15, UC-16, UC-17"],
        ],
        [1600, 3600, 4160],
    )

    doc.add_heading("4.4 Web pública: escaparate y entrada al chatbot", level=2)
    doc.add_paragraph(
        "La pantalla inicial concentra la identidad visual de Corte Perfecto, el acceso rápido a la reserva y un resumen "
        "de los servicios principales. La interfaz utiliza una estética oscura con acentos dorados, coherente con la "
        "percepción premium planteada desde el Capítulo 1. El botón principal abre el chatbot, por lo que la reserva no "
        "depende de que el cliente busque un formulario convencional."
    )
    add_picture(doc, SCREENSHOT_DIR / "01_home.png", "Figura 4.3. Pantalla de inicio de Corte Perfecto.")

    doc.add_paragraph(
        "La sección de servicios materializa el catálogo oficial de la peluquería. En la web se presentan los servicios "
        "base y sus precios; en el chatbot se utiliza el mismo catálogo en formato numerado del 1 al 7 para reducir "
        "ambigüedades conversacionales. Esta decisión conecta directamente con la corrección realizada tras las pruebas "
        "iniciales: la opción numérica evita que el modelo interprete mal expresiones como “el 4”."
    )
    add_picture(doc, SCREENSHOT_DIR / "02_servicios.png", "Figura 4.4. Sección de servicios con catálogo visible.")

    doc.add_heading("4.5 Chatbot de reserva con LM Studio", level=2)
    doc.add_paragraph(
        "El chatbot es el punto más singular del sistema. La interfaz se comporta como un asistente de reserva: saluda, "
        "ofrece información, acepta números de servicio, pregunta solo los datos que faltan y muestra una tarjeta cuando "
        "la cita se guarda correctamente. El componente `ChatWidget` mantiene el historial reciente y el identificador "
        "de conversación, pero la validación real se ejecuta en backend."
    )
    add_picture(doc, SCREENSHOT_DIR / "03_chat_abierto.png", "Figura 4.5. Widget de chat abierto sobre la web pública.")

    doc.add_paragraph(
        "La implementación evita confiar ciegamente en la salida del LLM. Primero se ejecutan reglas deterministas "
        "rápidas con `chatRuleService`; si el caso lo requiere, `promptService` prepara el historial y `lmStudioService` "
        "consulta LM Studio mediante endpoint compatible con OpenAI; finalmente `responseParserService` separa el mensaje "
        "del posible candidato de cita y `appointmentService` lo valida antes de persistir. Así se "
        "mantiene una conversación natural sin trasladar a la IA decisiones críticas de agenda."
    )
    add_picture(doc, booking_flow, "Figura 4.6. Flujo técnico de reserva por chatbot.")

    add_table(
        doc,
        ["Regla", "Lugar de implementación", "Efecto visible"],
        [
            ["Servicios por número", "serviceCatalog.js, chatRuleService.js, bookingFlowService.js", "El cliente puede responder 1..7 y el sistema resuelve el servicio exacto."],
            ["Nombre real obligatorio", "bookingFlowService.js, appointmentService.js", "No se registra una cita con nombre vacío, genérico o inválido."],
            ["Fin de semana cerrado", "calendarService.js, chatRuleService.js, bookingFlowService.js y appointmentService.js", "El sistema propone viernes o lunes antes de pedir hora."],
            ["Hora actual", "calendarService.js, chatRuleService.js y appointmentService.js", "No se aceptan reservas para una hora pasada del mismo día."],
            ["Solapes", "appointmentService.js", "No se crean dos citas activas en el mismo intervalo."],
            ["LM Studio caído", "lmStudioService.js, AppError.js y errorMiddleware.js", "El usuario recibe un error controlado y no se inventa una confirmación."],
        ],
        [1900, 3300, 3960],
    )

    doc.add_heading("4.6 Acceso y panel de administración", level=2)
    doc.add_paragraph(
        "El peluquero dispone de un panel privado. El acceso se realiza mediante usuario y contraseña; el backend valida "
        "la contraseña con hash y emite un JWT. Las rutas de citas quedan protegidas por middleware, de modo que un "
        "usuario sin token válido no puede consultar, modificar ni eliminar reservas."
    )
    add_picture(doc, SCREENSHOT_DIR / "04_admin_login.png", "Figura 4.7. Pantalla de inicio de sesión del administrador.")

    doc.add_paragraph(
        "Tras iniciar sesión, el dashboard resume la actividad: citas del día, pendientes, confirmadas, completadas, "
        "ingresos estimados y total de citas. Esta pantalla no sustituye a la agenda detallada, pero ofrece al peluquero "
        "una lectura rápida del estado del negocio."
    )
    add_picture(doc, SCREENSHOT_DIR / "05_admin_dashboard.png", "Figura 4.8. Dashboard administrativo.")

    doc.add_paragraph(
        "La gestión de citas permite filtrar por estado, acotar por fechas, ordenar cronológicamente y ejecutar acciones "
        "directas sobre cada reserva. Las acciones principales son editar, eliminar y marcar como completada. Esta última "
        "función fue incorporada para reflejar el ciclo de vida real de una cita: no basta con crearla, también debe "
        "poder cerrarse cuando el servicio ha sido prestado."
    )
    add_picture(doc, SCREENSHOT_DIR / "06_admin_citas.png", "Figura 4.9. Gestión administrativa de citas.")

    doc.add_paragraph(
        "La creación manual cubre el escenario en el que el cliente llama por teléfono, acude presencialmente o el "
        "peluquero necesita registrar una reserva sin pasar por el chatbot. El formulario utiliza el mismo backend y las "
        "mismas reglas de validación que las citas conversacionales, por lo que la agenda mantiene una única fuente de verdad."
    )
    add_picture(doc, SCREENSHOT_DIR / "07_admin_crear_cita.png", "Figura 4.10. Formulario de creación manual de cita.")

    doc.add_heading("4.7 Casos de uso representativos resueltos en interfaz", level=2)
    doc.add_paragraph(
        "La siguiente tabla recoge los casos de uso más representativos de la solución. La intención es mostrar el puente "
        "entre el modelado de capítulos anteriores y la experiencia implementada: cada acción del usuario tiene pantalla, "
        "endpoint, servicio y validación asociada."
    )
    add_table(
        doc,
        ["Caso de uso", "Interfaz", "Backend implicado", "Resultado"],
        [
            ["UC-02 Consultar servicios", "Sección Servicios y acción rápida del chat.", "GET /api/services, serviceCatalogService.", "Catálogo visible, numerado y consistente con MongoDB."],
            ["UC-05 Reservar cita por chatbot", "Widget de chat con conversación guiada.", "chatController, bookingFlowService, appointmentService.", "Cita confirmada y persistida en `appointments`."],
            ["UC-09 Modificar reserva activa", "Chat conserva la cita activa durante la conversación.", "updateAppointment con `activeAppointmentId`.", "La cita existente se actualiza sin duplicarla."],
            ["UC-10 Iniciar sesión", "Formulario de login administrador.", "authController, adminService, bcryptjs, JWT.", "Token válido y acceso al panel privado."],
            ["UC-12 Listar/filtrar citas", "Tabla administrativa con filtros y orden.", "listAppointments con query params.", "Agenda consultable por estado, fecha y orden."],
            ["UC-13 Crear cita manual", "Formulario de alta del panel.", "createAppointment con source admin.", "Reserva manual validada igual que las del chat."],
            ["UC-14 Editar cita", "Modal de edición sobre la tabla.", "updateAppointment.", "Cambio controlado de datos, servicio, estado o notas."],
            ["UC-15 Completar cita", "Botón de check en cada fila activa.", "PATCH /api/appointments/:id.", "La cita pasa a estado `completed`."],
            ["UC-16 Eliminar cita", "Botón de papelera con confirmación.", "deleteAppointment.", "El registro se elimina de MongoDB."],
        ],
        [2100, 2500, 2500, 2260],
    )

    doc.add_heading("4.8 Persistencia y sincronización de datos", level=2)
    doc.add_paragraph(
        "La base de datos `corte_perfecto` utiliza un modelo documental sencillo y suficiente para la primera versión "
        "del sistema. Las citas se guardan como documentos completos con nombre, servicio, precio, duración, fecha, hora, "
        "rango temporal, estado, origen y notas. El catálogo de servicios se sincroniza al arrancar desde una definición "
        "oficial del backend, lo que permite que MongoDB refleje las siete opciones disponibles sin convertir el catálogo "
        "en lógica duplicada."
    )
    add_table(
        doc,
        ["Colección", "Uso", "Campos principales"],
        [
            ["appointments", "Agenda operativa de reservas.", "customerName, service, price, duration, date, time, startsAt, endsAt, status, source, notes, conversationId."],
            ["admins", "Cuenta de acceso del peluquero.", "username, passwordHash, role, timestamps."],
            ["servicios", "Catálogo público sincronizado.", "id, key, nombre, descripcion, precio, duracion_minutos."],
        ],
        [1800, 2800, 4760],
    )

    doc.add_heading("4.9 Verificación de la implementación", level=2)
    doc.add_paragraph(
        "La verificación combina pruebas automáticas y comprobación manual de la interfaz. Siguiendo la filosofía de "
        "pySigHor, la evidencia no queda únicamente en el texto del TFG: también permanece en el repositorio mediante "
        "la carpeta `RUP/99-seguimiento` y las pruebas del backend."
    )
    add_table(
        doc,
        ["Comprobación", "Comando / artefacto", "Estado"],
        [
            ["Sintaxis backend", "npm run check --prefix backend", "Superado"],
            ["Reglas de negocio", "npm run test --prefix backend", "10 pruebas superadas"],
            ["Build frontend", "npm run build --prefix frontend", "Superado"],
            ["Verificación integrada", "npm run verify", "Superado"],
            ["Trazabilidad UC-código", "RUP/99-seguimiento/trazabilidad-casos-uso.md", "Actualizada"],
            ["Auditoría diseño-implementación", "RUP/99-seguimiento/auditoria-diseno-implementacion.md", "Actualizada"],
        ],
        [2600, 4200, 2560],
    )

    doc.add_heading("4.10 Recorrido end-to-end de una reserva", level=2)
    doc.add_paragraph(
        "El flujo completo de reserva confirma que la solución no funciona como piezas aisladas, sino como un circuito "
        "integrado. El cliente inicia la conversación en la web pública; el frontend envía el mensaje al backend; el "
        "backend decide si puede responder mediante reglas propias o si necesita consultar LM Studio; cuando se reúnen "
        "nombre, servicio, fecha y hora, la cita se valida y se guarda en MongoDB. Finalmente, el frontend muestra una "
        "tarjeta de confirmación y el panel administrativo puede consultar la misma reserva."
    )
    add_table(
        doc,
        ["Paso", "Elemento responsable", "Control aplicado"],
        [
            ["1. Inicio de conversación", "ChatWidget.jsx", "Crea identificador de conversación y conserva historial reciente."],
            ["2. Entrada al backend", "chatController.js", "Rechaza mensajes vacíos y coordina el flujo."],
            ["3. Reglas previas", "chatRuleService.js", "Responde a horarios/servicios, detecta fines de semana y opciones numéricas inválidas."],
            ["4. Flujo de reserva", "bookingFlowService.js", "Extrae nombre, servicio, fecha y hora manteniendo contexto conversacional."],
            ["5. IA local", "lmStudioService.js", "Solo se consulta cuando hace falta conversación abierta; timeout y error controlado."],
            ["6. Persistencia", "appointmentService.js y Appointment.js", "Valida nombre, servicio, horario, fecha futura y solapes antes de guardar."],
            ["7. Confirmación", "ChatWidget.jsx", "Muestra respuesta natural y tarjeta con datos persistidos."],
            ["8. Gestión posterior", "AdminAppointments.jsx", "Permite editar, completar o eliminar la misma cita."],
        ],
        [900, 2700, 5760],
    )

    doc.add_heading("4.11 Revisión MVC y principios de diseño", level=2)
    doc.add_paragraph(
        "El código no implementa SOLID de forma ceremonial con clases innecesarias; lo aplica en decisiones concretas: "
        "responsabilidades pequeñas, dependencias claras y servicios especializados. En un proyecto Node.js moderno, "
        "muchas unidades de diseño son módulos exportados en lugar de clases ES6, pero la separación conceptual se "
        "mantiene."
    )
    add_table(
        doc,
        ["Principio", "Aplicación en Corte Perfecto", "Ejemplo"],
        [
            ["Responsabilidad única", "Cada servicio concentra una familia de reglas.", "calendarService para fechas; appointmentService para agenda; promptService para prompt."],
            ["Abierto/cerrado", "El catálogo se amplía desde una fuente central sin tocar controladores.", "SERVICE_CATALOG y serviceCatalogService."],
            ["Sustitución", "La API de LM Studio queda encapsulada, permitiendo cambiar el proveedor compatible.", "lmStudioService."],
            ["Segregación de interfaces", "El frontend consume fachadas pequeñas por dominio.", "authApi, appointmentApi, chatApi y serviceApi."],
            ["Inversión de dependencias práctica", "Controladores dependen de servicios, no de detalles de MongoDB.", "appointmentController delega en appointmentService."],
            ["MVC", "Rutas y controladores son entrada; servicios son negocio; modelos son persistencia; React contiene vistas.", "backend/src y frontend/src."],
        ],
        [1900, 4200, 3260],
    )

    doc.add_heading("4.12 Cierre del capítulo", level=2)
    doc.add_paragraph(
        "La solución implementada cubre el flujo completo planteado al inicio del trabajo: un cliente puede conocer los "
        "servicios, conversar con un asistente local, reservar una cita y recibir confirmación; el peluquero puede "
        "autenticarse, consultar la agenda, crear citas manualmente, modificarlas, completarlas o eliminarlas. La "
        "implementación conserva la separación de responsabilidades definida en el diseño y evita que la IA tenga control "
        "directo sobre la persistencia o sobre las reglas críticas de agenda."
    )

    doc.add_page_break()
    doc.add_heading("Capítulo 5. Conclusiones y líneas futuras", level=1)
    doc.add_paragraph(
        "El último capítulo valora el resultado obtenido. A diferencia del Capítulo 1, centrado en presentar el problema, "
        "este capítulo contrasta la hipótesis y los objetivos con la evidencia generada durante el desarrollo. También "
        "recoge las decisiones que han funcionado, las limitaciones razonables de una primera versión y las líneas de "
        "continuidad que permitirían evolucionar la plataforma."
    )

    doc.add_heading("5.1 Cumplimiento del objetivo general", level=2)
    doc.add_paragraph(
        "El objetivo general era diseñar e implementar una plataforma web de gestión de reservas para Corte Perfecto que "
        "automatizara la atención al cliente mediante un asistente conversacional local, persistiera las reservas en "
        "MongoDB y proporcionara una interfaz de administración. El resultado cumple ese objetivo: existe una web pública "
        "funcional, un chatbot conectado a LM Studio, una API Express con reglas de negocio, una base de datos MongoDB y "
        "un panel privado para el peluquero."
    )
    add_callout(
        doc,
        "Conclusión principal",
        "La hipótesis de partida se considera validada dentro del alcance del TFG: una peluquería pequeña puede automatizar parte de su reserva de citas con IA local, sin enviar datos personales a proveedores externos de inferencia y sin pagar por cada conversación.",
    )

    doc.add_heading("5.2 Cumplimiento de objetivos específicos", level=2)
    add_table(
        doc,
        ["Objetivo", "Evidencia de cumplimiento", "Valoración"],
        [
            ["OE1 Requisitos", "Capítulo 2 define actores, casos de uso, modelo de dominio, estados, contexto y trazabilidad RS-UC.", "Cumplido"],
            ["OE2 Análisis y diseño", "Capítulo 3 concreta arquitectura, MVC, paquetes, modelo documental, secuencias, prompt y contingencia IA.", "Cumplido"],
            ["OE3 Producto funcional", "Capítulo 4 muestra pantallas reales, API, MongoDB, chatbot LM Studio y panel administrador.", "Cumplido"],
            ["OE4 Evaluación", "Pruebas automáticas, `npm run verify`, auditoría diseño-código y revisión de reglas críticas.", "Cumplido"],
        ],
        [2300, 5100, 1960],
    )

    doc.add_heading("5.3 Cumplimiento de objetivos transversales", level=2)
    add_table(
        doc,
        ["Objetivo transversal", "Resultado obtenido"],
        [
            ["OET1 Privacidad", "La inferencia se ejecuta en LM Studio local. Los datos de reserva no se envían a una API externa de IA; se almacenan en MongoDB local."],
            ["OET2 Robustez del chatbot", "El flujo combina prompt estructurado, catálogo numerado, reglas previas, parser de respuesta y validaciones de backend."],
            ["OET3 Usabilidad", "El cliente puede reservar conversando y el administrador dispone de acciones directas para gestionar la agenda."],
            ["OET4 Mantenibilidad", "La separación frontend/backend/servicios/modelos permite modificar reglas de agenda sin reescribir la interfaz."],
        ],
        [2600, 6760],
    )

    doc.add_heading("5.4 Evaluación de eficiencia e integridad", level=2)
    doc.add_paragraph(
        "La eficiencia del sistema se evalúa desde dos perspectivas. La primera es técnica: el backend debe validar y "
        "persistir citas sin inconsistencias. La segunda es operativa: el usuario debe poder completar una reserva sin "
        "navegar por formularios largos. El tiempo exacto de generación del LLM depende del equipo local y del modelo "
        "cargado en LM Studio, por lo que la aplicación no fija una cifra universal; en su lugar incorpora timeout "
        "configurable, endpoint de salud y respuesta de error controlada."
    )
    add_table(
        doc,
        ["Aspecto evaluado", "Mecanismo", "Resultado"],
        [
            ["Integridad de citas", "Validación de nombre, servicio, horario, fecha futura y solape.", "No se persiste una reserva inválida."],
            ["Coherencia de catálogo", "Sincronización de `servicios` desde SERVICE_CATALOG.", "Los servicios 1..7 son consistentes en prompt, API y MongoDB."],
            ["Disponibilidad de IA", "Timeout `LMSTUDIO_TIMEOUT_MS` y endpoint `/api/health/lmstudio`.", "Fallo controlado si LM Studio no responde."],
            ["Rendimiento frontend", "Build Vite de producción.", "Aplicación preparada para servir assets optimizados."],
            ["Regresión funcional", "`node:test` sobre servicios críticos.", "10 pruebas automáticas superadas."],
            ["Operación administrativa", "Filtros, orden y estados en panel.", "El peluquero puede consultar y cerrar el ciclo de vida de la cita."],
        ],
        [2400, 3800, 3160],
    )

    doc.add_heading("5.5 Discusión de resultados", level=2)
    doc.add_paragraph(
        "La decisión más relevante del proyecto ha sido separar conversación y decisión. LM Studio aporta naturalidad, "
        "tono y flexibilidad lingüística; sin embargo, el sistema no le concede autoridad final sobre la reserva. Esta "
        "arquitectura reduce el riesgo de alucinaciones operativas: aunque el modelo produzca texto incorrecto, la cita "
        "solo se registra si supera las reglas del backend."
    )
    doc.add_paragraph(
        "Otra decisión importante ha sido utilizar MongoDB. En una agenda de citas, cada reserva se consulta como unidad "
        "documental completa y no requiere un modelo relacional complejo. El esquema usado conserva los datos necesarios "
        "para el panel, las estadísticas y la detección de solapes. Además, Mongoose aporta validación de esquema, índices "
        "y una capa de acceso coherente con Node.js."
    )
    doc.add_paragraph(
        "Desde el punto de vista metodológico, la inspiración en pySigHor ha sido especialmente útil para no dejar el TFG "
        "como una simple implementación. La carpeta RUP actúa como memoria viva del proyecto: casos de uso, código, "
        "pruebas y auditoría pueden recorrerse juntos. Esta trazabilidad facilita justificar decisiones ante el tutor y "
        "ante un tribunal, porque cada pantalla se conecta con un caso de uso y cada regla crítica con una prueba."
    )

    add_table(
        doc,
        ["Decisión", "Ventaja", "Compromiso asumido"],
        [
            ["IA local con LM Studio", "Privacidad y ausencia de coste por llamada.", "Dependencia de que el equipo local tenga el modelo cargado."],
            ["Reglas deterministas en backend", "Mayor seguridad funcional frente a respuestas impredecibles.", "Más lógica propia que mantener."],
            ["React/Vite", "Interfaz rápida de desarrollar, modular y moderna.", "Necesidad de build separado para producción."],
            ["Node.js/Express", "API ligera, comprensible y suficiente para MVC.", "Menos estructura impuesta que frameworks más opinados."],
            ["MongoDB/Mongoose", "Modelo documental simple y flexible.", "Requiere cuidar índices y validaciones para mantener consistencia."],
        ],
        [2200, 3600, 3560],
    )

    doc.add_heading("5.6 Limitaciones detectadas", level=2)
    doc.add_paragraph(
        "Las limitaciones no invalidan la solución; delimitan el alcance realista de una primera versión de TFG. La más "
        "evidente es que el chatbot depende del servidor local de LM Studio. Si el modelo no está cargado o el puerto "
        "no responde, el sistema informa del problema y evita confirmar reservas inventadas, pero la experiencia "
        "conversacional queda temporalmente indisponible."
    )
    add_bullets(
        doc,
        [
            "La disponibilidad del asistente depende de que LM Studio esté abierto, el modelo cargado y el endpoint local activo.",
            "La agenda no incorpora aún notificaciones automáticas por correo, SMS o WhatsApp.",
            "El sistema está parametrizado para una única peluquería y no para una red de establecimientos.",
            "La autenticación cubre al administrador, pero no existe todavía un área privada para clientes finales.",
            "La evaluación se centra en pruebas funcionales y de reglas de negocio; no se ha realizado un estudio formal con usuarios reales.",
        ],
    )

    doc.add_heading("5.7 Recomendaciones", level=2)
    doc.add_paragraph(
        "Para una puesta en uso real conviene mantener el mismo criterio que ha guiado el desarrollo: avanzar por "
        "iteraciones pequeñas, con trazabilidad y pruebas. No sería recomendable añadir funcionalidades de forma masiva "
        "sin consolidar primero la operación diaria del peluquero."
    )
    add_numbered(
        doc,
        [
            "Mantener el backend como autoridad de negocio: ningún cambio del prompt debe sustituir validaciones de agenda.",
            "Configurar variables de entorno propias antes de un despliegue real, especialmente `JWT_SECRET`, usuario administrador y URI de MongoDB.",
            "Añadir copias de seguridad periódicas de MongoDB si la aplicación se usa con clientes reales.",
            "Registrar conversaciones solo si existe consentimiento y una política clara de conservación de datos.",
            "Ampliar las pruebas automáticas cada vez que se añada una regla de negocio o un nuevo estado de cita.",
        ],
    )

    doc.add_heading("5.8 Futuras líneas de actuación", level=2)
    doc.add_paragraph(
        "La solución queda preparada para evolucionar. La estructura MVC, el catálogo centralizado y la separación entre "
        "IA, negocio y persistencia permiten añadir funcionalidades sin reescribir el núcleo."
    )
    add_table(
        doc,
        ["Línea futura", "Descripción", "Prioridad"],
        [
            ["Disponibilidad avanzada", "Mostrar huecos libres calculados automáticamente y sugerir alternativas concretas.", "Alta"],
            ["Recordatorios", "Enviar avisos por email, SMS o WhatsApp antes de la cita.", "Alta"],
            ["Cancelación por cliente", "Permitir cancelar o cambiar una reserva mediante enlace seguro o conversación.", "Media"],
            ["Calendario visual", "Añadir vista semanal/mensual para el peluquero.", "Media"],
            ["Multiusuario", "Soportar varios peluqueros, turnos y servicios por empleado.", "Media"],
            ["Analítica", "Informes de servicios más solicitados, horas punta e ingresos por periodo.", "Media"],
            ["Mejora del chatbot", "Añadir pruebas conversacionales de regresión y evaluación de calidad de respuesta.", "Alta"],
            ["Despliegue controlado", "Preparar ejecución en red local o servidor privado manteniendo IA local.", "Baja-media"],
        ],
        [2200, 5200, 1960],
    )

    doc.add_heading("5.9 Valoración personal del proceso", level=2)
    doc.add_paragraph(
        "El proyecto ha permitido recorrer un ciclo completo de ingeniería de software: entender una necesidad real, "
        "modelarla, diseñar una arquitectura, implementar una solución y verificar sus reglas críticas. La parte más "
        "valiosa no ha sido únicamente integrar un LLM, sino aprender a integrarlo con responsabilidad: la IA mejora la "
        "experiencia de usuario, pero no reemplaza el diseño de dominio ni la validación de negocio."
    )
    doc.add_paragraph(
        "También se confirma la utilidad de documentar el proceso. Los capítulos, los diagramas, la carpeta RUP y las "
        "pruebas no son elementos aislados; forman una cadena de evidencia. Esa cadena es la que permite defender que "
        "la aplicación no se ha construido de forma improvisada, sino siguiendo un razonamiento técnico progresivo."
    )

    doc.add_heading("5.10 Conclusión final", level=2)
    doc.add_paragraph(
        "Corte Perfecto demuestra que una pequeña empresa puede beneficiarse de tecnologías actuales sin asumir una "
        "arquitectura desproporcionada ni renunciar a la privacidad. La combinación de React, Node.js, MongoDB y LM Studio "
        "permite construir una solución funcional, local, mantenible y alineada con el RGPD. El resultado final satisface "
        "los objetivos planteados y deja una base técnica suficiente para seguir evolucionando el sistema en iteraciones "
        "posteriores."
    )

    doc.add_page_break()
    doc.add_heading("Anexo A. Evidencia de verificación final", level=1)
    doc.add_paragraph(
        "Este anexo resume las evidencias operativas utilizadas para cerrar la entrega. Se incluye para facilitar la "
        "revisión posterior, aunque los artefactos completos permanecen en el repositorio."
    )
    add_table(
        doc,
        ["Artefacto", "Ubicación", "Uso"],
        [
            ["Pruebas de calendario", "backend/tests/calendarService.test.js", "Validar días laborables, fines de semana y formato de fecha."],
            ["Pruebas de flujo conversacional", "backend/tests/bookingFlowService.test.js", "Validar nombre, selección numérica de servicios y bloqueo de fin de semana."],
            ["Pruebas de agenda", "backend/tests/appointmentService.test.js", "Validar citas, solapes, estados y catálogo sincronizado."],
            ["Matriz de trazabilidad", "RUP/99-seguimiento/trazabilidad-casos-uso.md", "Relacionar casos de uso con implementación y pruebas."],
            ["Auditoría diseño-código", "RUP/99-seguimiento/auditoria-diseno-implementacion.md", "Comprobar coherencia entre Capítulo 3 y código real."],
            ["Dashboard RUP", "RUP/99-seguimiento/estado-casos-uso.puml", "Visualizar el estado de los casos de uso."],
        ],
        [2300, 3600, 3460],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
